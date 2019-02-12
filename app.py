from docx import Document
import csv

#CONFIG SETTINGS
MODULE_CODE = "CMP5327";
FEEDBACK_TEMPLATE = 'config/coversheet.docx';
FEEDBACK_SPREADSHEET = 'config/feedback.csv'

def replaceTextInFeedbackSheet():
    document = Document(FEEDBACK_TEMPLATE)
    #sectionid increments when a section is written to
    sectionid = 0;
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for s in SECTION_IDENTIFIERS:
                        if s in p.text:
                            line = p.text;
                            line = line.replace(s,section_replacement[sectionid]);
                            sectionid+=1;#increments when a section is written to
                            p.text = line;
                            document.save("release/"+MODULE_CODE+" - Feedback coversheet " +section_replacement[0]+ ".docx")

SECTION_IDENTIFIERS = ["S1001", "S1002","S1003","S1004","S1005","S1006", "S1007","S1008", "S1009"];
section_replacement = []; #student_name, student_number, feedback_point_1, feedback_point_2, feedback_point_3, feedback_point_4, feedforward, grade, grade


feedback_document = open(FEEDBACK_SPREADSHEET);
reader = csv.reader(feedback_document);

#loop through rows and skip first
for row in reader:
    i = 0;
    for col in row:
        section_replacement.append(col);
        i+=1;
    replaceTextInFeedbackSheet();
    del section_replacement[:]
        #print(col);
